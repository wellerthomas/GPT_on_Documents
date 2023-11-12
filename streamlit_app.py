from langchain.chat_models import ChatOpenAI
from langchain.schema import (
    SystemMessage,
    HumanMessage,
    AIMessage
)

from langchain import PromptTemplate
from langchain.chains import LLMChain

import streamlit as st
from streamlit_chat import message
import tiktoken

import time
import re
import os
import pandas as pd
from docx import Document
from io import BytesIO



OPENAI_API_KEY=st.secrets['OPENAI_API_KEY']


def load_document(file):
    import os
    name, extension = os.path.splitext(file)

    if extension == '.pdf':
        from langchain.document_loaders import PyPDFLoader
        print(f'Loading {file}')
        loader  = PyPDFLoader(file)
    elif extension == '.docx':
        from langchain.document_loaders import Docx2txtLoader
        print(f'Loading {file}')
        loader  = Docx2txtLoader(file)
    elif extension == '.csv':
        from langchain.document_loaders import CSVLoader
        print(f'Loading {file}')
        loader  = CSVLoader(file)
    elif extension == '.txt':
        from langchain.document_loaders import TextLoader
        loader  = TextLoader(file)
    else:
        print('Document format is not supported!')
        return None 

    data = loader.load()
    return data

def num_tokens_from_string(string: str, encoding_name: str) -> int:
    encoding = tiktoken.encoding_for_model(encoding_name)
    num_tokens = len(encoding.encode(string))
    return num_tokens

def create_document(output):
    document = Document()
    document.add_heading('Summary of the Document:', level=1)

    # Regular expression to find **Headline Text**
    headline_regex = r'\*\*(.*?)\*\*'

    # Split the text by lines
    lines = output.split('\n')

    for line in lines:
        # Search for headlines
        headline_search = re.search(headline_regex, line)
        if headline_search:
            # If a headline is found, add it as a heading
            headline_text = headline_search.group(1).strip()
            document.add_heading(headline_text, level=2)
            continue  # Skip adding this line as a paragraph

        # Add non-headline lines as paragraphs
        if line.strip() != '':  # Skip empty lines
            document.add_paragraph(line)

    return document

st.set_page_config(
    page_title='GPT on Documents',
    page_icon=':memo:'
)
st.title('GPT on Documents')
st.subheader('''Upload the files''')

uploaded_files = st.file_uploader("Upload files:", type=['pdf', 'docx', 'csv', 'txt'], accept_multiple_files=True)

# Load data
file_inputs = []
if uploaded_files:
    for uploaded_file in uploaded_files:
        bytes_data = uploaded_file.read()
        file_name = os.path.join('./', uploaded_file.name)
        with open(file_name, 'wb') as f:
            f.write(bytes_data)

        file_input = load_document(file_name)
        file_inputs.append(file_input)


    # Initialize an empty string to hold all the text content
    all_file_data = ''

    # Loop over each file's list in file_inputs (neccessary for the number of token)
    for file_list in file_inputs:
        # Each file_list contains Document objects for each page
        for document in file_list:
            # Extract the text content from the Document object's page_content attribute
            text_content = document.page_content
            # Append the text content to the all_file_data string
            all_file_data += text_content + '\n'  # Adding a newline as a page separator

    # Now all_file_data contains the concatenated text from all the pages of all the Document objects



    # Initialize session state for user_prompt
    if 'user_prompt' not in st.session_state:
        st.session_state['user_prompt'] = "Summarize the information."


    # Use the session state variable for the text area, and update it on user input
    st.session_state['user_prompt'] = st.text_area(
        label='Your Prompt', 
        value=st.session_state['user_prompt'], 
        height=100
    )


    template_string = '''
    There are one or multiple documents loaded, which are represented by the variable {file_inputs}. Based on these documents, here is the task: {user_prompt}

    Please ensure the output is structured suitably for conversion into a Word document, with key points highlighted as headlines. Headlines should be marked with ** at the start and end for emphasis.
    '''

    prompt_data = PromptTemplate(
        input_variables=['file_inputs', 'user_prompt'],
        template=template_string
    )

    st.markdown("""
    <style>
    .stButton button {
        background-color: #0074D9;
        color: #FFFFFF;
    }
    </style>
    """, unsafe_allow_html=True)

    # Add the "Start" button
    if st.button('Start'):


        # Make sure to pass both file_inputs and user_prompt
        chain_input = {
            'file_inputs': file_inputs,
            'user_prompt': st.session_state['user_prompt']
        }

        llm = ChatOpenAI(temperature=0, model_name='gpt-3.5-turbo')
        #llm = ChatOpenAI(temperature=0, model_name='gpt-4')

        number_of_token = num_tokens_from_string(all_file_data, "gpt-3.5-turbo")

        # Check if the number of tokens exceeds the maximum limit
        if number_of_token > 12000:
            st.error("The length of the document is too long for processing. Please upload a shorter document.")
            st.stop()
        elif number_of_token > 3000:
            # If more than 3000 tokens, switch to the larger model
            llm = ChatOpenAI(temperature=0, model_name="gpt-3.5-turbo-16k")
            processing_text = st.text("In progress... and switch to gpt-3.5-turbo-16k because of the document sizes.")
        else:
            # If 3000 tokens or fewer, proceed with the original model
            processing_text = st.text("One moment, I'm working on that...")

        chain = LLMChain(llm=llm, prompt=prompt_data)
        output_add_infos = chain.run(chain_input)

        output_add_infos_cleaned = output_add_infos.replace('**', '')

        processing_text.empty()
        
        # Create the document
        doc = create_document(output_add_infos)

        # Save the document to a BytesIO object
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        # Use the download button to offer the document for download
        st.download_button(
            label="Download Word document",
            data=buffer,
            file_name="output.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

        st.markdown(output_add_infos, unsafe_allow_html=True)
        
else:
    # Display a message if no files have been uploaded
    st.warning('Please upload at least one document to start.')
